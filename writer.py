from docx import Document
from docx.oxml.ns import qn
from docx.shared import  Pt
from docx.enum.style import WD_STYLE_TYPE
import numpy as np
class Writers():
    def __init__(self, reader, norm,test):
        self.norm = norm
        self.reader = reader
        self.doc = Document()
        self.test=test
        self.doc.styles['Normal'].font.name = u'標楷體'
        self.doc.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'標楷體')
        self.doc.styles['Normal'].font.bold = True
        self.doc.styles['Normal'].font.size = Pt(12)
        styles = self.doc.styles
        new_heading_style = styles.add_style(
            'New Heading', WD_STYLE_TYPE.PARAGRAPH)

        font = new_heading_style.font
        font.bold = True
        font.name = u'標楷體'
        font.size = Pt(16)
        new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        self.tables = []
        tables = self.tables
        title = ['']
        i = 0
        if test=='軌跡標示測驗':
            self.add_paragraph('')
            self.add_title('軌跡標示測驗')
            self.add_paragraph('基本測量:')
            table1 = Table( [['','情境一：\n視覺掃描','情境二：\n圓形序列', '情境三：\n六邊形序列','情境四：\n圓形六邊形轉換','情境五：\n動作速度']],
            test, i, self.doc,norm, reader)
            tables.append(table1)
            i += 1
            self.add_paragraph('')
            self.add_paragraph('衍生測量:')
            getTable = [x[1:] for x in table1.tableList[1:]]
            table2 = Table([['','圓形序列＋六邊形序列','圓形六邊形轉換 - 視覺掃描','圓形六邊形轉換 - 圓形序列','圓形六邊形轉換 - 六邊形序列','圓形六邊形轉換- 圓形序列＋六邊形序列','圓形六邊形轉換 - 動作速度']],
            test, i, self.doc,norm, getTable)
            tables.append(table2)
        elif test=="設計流暢測驗":
            self.add_paragraph('')
            self.add_title('設計流暢性測驗')
            self.add_paragraph('基本測量:')
            table1 = Table( [['','總正確數'],['','情境1\n黑點相連','情境2\n白點相連',       '情境3 黑點和白點互相轉換','設計流暢性\n總和']],
            test, i, self.doc,norm, reader,merge=[(0,1),(0,4)])
            tables.append(table1)
            i += 1
            self.add_paragraph('')
            self.add_paragraph('衍生測量:')
            getTable = [x[1:] for x in table1.tableList[2:]]
            table2 = Table([['',  '實心點連結＋空心點連結\n組合總正確數','轉換 - 實心點連結＋空心點連結']],
            test, i, self.doc,norm, getTable)
            tables.append(table2)
            i += 1
            self.add_paragraph('')
            self.add_paragraph('選擇性測量')
            getTable = [x[1:] for x in table1.tableList[2:]]
            table2 = Table([['',        '不正確設計數','重複設計','嘗試設計總數','正確設計\n百分比']],
            test, i, self.doc,norm,reader,getTable)
            tables.append(table2)
        '''
        i += 1
        getTable = [x[1:] for x in table1.tableList[2:]]
        table3 = Table(title, 'Track', i, self.doc,self.getNormReverse()[2],norm, getTable)
        tables.append(table3)
        i += 1
        '''


    

    def add_title(self, title):
        paragraph = self.doc.add_paragraph(title, style='New Heading')
        paragraph.alignment = 1

    def add_paragraph(self, context):
        self.doc.add_paragraph(context)

    def save(self, path):
        self.doc.save(path)

    def write_all(self):
        pass
class OverwriteError(Exception):
    def __init__(self,method1):
        super(Exception, self).__init__('This method should be overwritten', method1+' is not overwritten')
        self.cause = cause
        self.message = message
    def __str__(self):
        return self.cause + ': ' + self.message
class Table_Track(Table):
    def getNormReverse(self):
        talbe2 = [True, True, True, True, True]
        talbe3 = [False, False, False, False, False, False]
        tables = [talbe2, talbe3]
        return tables[self.table]
class Table_Fluent(Table):
    def getNormReverse(self):
        table2=[False,False,False,False]
        table3=[False,False]
        table4=[True,True,False,False]
        tables=[table2,table3,table4]
        return tables[self.table]

class Table():
    def __init(self, title, word, table, doc,norm, *args,merge=[]):
        self.doc = doc
        self.word=word
        self.table=table
        self.norm=norm
        reverse=self.getNormReverse()
        print('reverse',reverse)
        i1 = title
        print('i2 args:',len(args))
        i2 = ['原始\n分數']+self.getData(*args)
        print(i2)
        i3 = ['量尺\n分數']+self.getScaleAndPr('table'+str(2+table),
                                           i2[1:], reverse=reverse)[0]
        i4 = ['PR值']+self.getScaleAndPr('table'+str(2+table),
                                        i2[1:], reverse=reverse)[1]
        self.tableList = i1+[ i2, i3, i4]
        self.add_table([len(self.tableList), len(i2)], self.tableList,merge)
    def getNormReverse(self):
        raise OverwriteError('getNormReverse')

    def getData(self, *args):
        word=self.word
        table=self.table
        if word == '軌跡標示測驗':
            if table == 0:
                reader = args[0]
                print('reader',[reader.getData('Task1', 'Complete_Time'), reader.getData('Task2', 'Complete_Time'), reader.getData('Task3', 'Complete_Time'), reader.getData('Task4', 'Complete_Time'), reader.getData('Task5', 'Complete_Time')])
                return [reader.getData('Task1', 'Complete_Time'), reader.getData('Task2', 'Complete_Time'), reader.getData('Task3', 'Complete_Time'), reader.getData('Task4', 'Complete_Time'), reader.getData('Task5', 'Complete_Time')]
            elif table == 1:
                table2 = args[0]
                table2 = table2[1]
                return [table2[1]+table2[2],   table2[0]-table2[1],        table2[3]-table2[1],      table2[3]-table2[2],       table2[3]-table2[1]-table2[2],                table2[3]-table2[4]]
        elif word == "設計流暢測驗":
            if table == 0:
                reader = args[0]
                getDataArg = ('Examing', 'Total_CorrectDesign', False)
                i2 = [int(reader.getData(*getDataArg, '情境1_黑點相連')),                           int(reader.getData(
                    *getDataArg, '情境2_白點相連')),          int(reader.getData(*getDataArg, '情境3_黑點和白點互相轉換'))]
                i2end = [sum(i2)]
                i2 = i2+i2end
                return i2
            elif table == 1:
                table2 = args[0]
                table2 = table2[1]
                return [table2[0]+table2[1],   table2[2]-table2[0]-table2[1]]
            elif table == 2:
                reader, table2 = args
                row = table2[0]
                getData = reader.getData
                getDataArg1 = ['Examing', 'Total_UnCorrectDesign', False]
                getDataArg2 = ['Examing', 'Total_RepeatDesign', False]
                getDataArg3 = ['Examing', 'Total_TryDesign', False]
                final = [getData(*getDataArg1, '情境1_黑點相連')+getData(*getDataArg1, '情境2_白點相連')+getData(*getDataArg1, '情境3_黑點和白點互相轉換'),
                getData(*getDataArg2, '情境1_黑點相連')+getData(*getDataArg2,
                        '情境2_白點相連')+getData(*getDataArg2, '情境3_黑點和白點互相轉換'),
                getData(*getDataArg3, '情境1_黑點相連')+getData(*getDataArg3, '情境2_白點相連')+getData(*getDataArg3, '情境3_黑點和白點互相轉換')]
                print('final',final)
                print('row',row)
                finalEnd = [row[3]/final[2]]
                return final+finalEnd
    def getScaleAndPr(self, table, data, reverse):
        scale = []
        norm = self.norm[table]
        # print('normmmmm',norm)
        for index, column in enumerate(norm):
            for itemIndex, item in enumerate(column):
                if (reverse[index] and data[index] >= item[0]) or (not reverse[index] and data[index] <= item[0]):
                    if data[index] == item[0]:
                        scale.append(item[1:])
                    else:
                        if itemIndex == 0:
                            scale.append(item[1:])
                        else:
                            result = []
                            scaleResult = self.getInnerValue(
                                data[index], column[itemIndex-1][0], column[itemIndex][0], column[itemIndex-1][1], column[itemIndex][1])
                            result = result+[scaleResult]
                            cumulativePercentage = self.getInnerValue(
                                data[index], column[itemIndex-1][0], column[itemIndex][0], column[itemIndex-1][3], column[itemIndex][3])
                            pr = int(cumulativePercentage*100)
                            if pr == 0:
                                pr = 1
                            result = result+[pr]
                            result = result+[cumulativePercentage]
                            scale.append(result)
                    break
                elif itemIndex == len(column)-1:  # if you didn't catch anything
                    scale.append(item[1:])
        scale=np.array(scale)
        scale=scale.transpose().tolist()
        print('scale',scale)
        scale[1]=list(map(int,scale[1]))
        return scale
    def add_table(self, size, context, merge=[]):
        table = self.doc.add_table(
            rows=size[0], cols=size[1], style='TableGrid')
        for index, item in enumerate(context):
            for index2, item2 in enumerate(item):
                if isinstance(item2,float):
                    table.cell(index,index2).text=f'{item2:.3f}'
                else:
                    table.cell(index,index2).text=str(item2)
                
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = 1
        if merge:
            a=table.cell(*merge[0])
            b=table.cell(*merge[1])
            a.merge(b)





