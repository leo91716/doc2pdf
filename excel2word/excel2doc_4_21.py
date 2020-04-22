from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import xlrd, datetime
from docx.shared import  Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import comtypes.client
from scipy.stats import norm
import glob
from tkinter import *
from tkinter.ttk import *
from tkinter import messagebox
from tkinter import filedialog

class Reader():
    def __init__(self, source):
        book = xlrd.open_workbook(source)
        #print('book', type(book))
        self.book=book
        sh = book.sheet_by_index(0)
        self.sh=sh

    def getData(self, data, task, sh=None):
        if sh==None:
            sh=self.sh
        for count in range(sh.nrows):
            if sh.cell(count, 0).value==task:
                while  True:
                    if sh.cell(count, 1).value=='Complete_Time':
                        data[task]=sh.cell(count, 2).value/1000
                        break
                    count+=1
                
                break

    def getBasicInfo(self, collect1, sh=None,book=None):
        if sh==None:
            sh=self.sh
        if book==None:
            book=self.book
        for count in range(sh.nrows):
            if sh.cell(count, 0).value=='個人資料':
                collect1[sh.cell(count, 1).value]=sh.cell(count, 2).value
                count+=1
                while  sh.cell(count, 0).value=='':
                    if sh.cell(count, 1).value=='DateOfBirth':
                        collect1[sh.cell(count, 1).value]=datetime.datetime(*xlrd.xldate_as_tuple(sh.cell(count, 2).value, book.datemode))
                    else:
                        collect1[sh.cell(count, 1).value]=sh.cell(count, 2).value
                    count+=1
                break

        #print(collect1)
        collect1[sh.cell(0, 3).value]=datetime.datetime.strptime(sh.cell(1, 3).value,'%Y/%m/%d-%H:%M:%S-%f')





def prValue(data,mean,std, norm=norm):
    if norm.cdf(data,mean,std)*100<1:
        return str(1)
    else:
        return str(int(norm.cdf(data,mean,std)*100))









def newScore(data, mean,std, new_mean, new_std):
    return (data-mean)/std*new_std+new_mean



def excel2doc(reader, file, destFolder):
    #print(datetime.datetime.strptime(sh.cell(1, 3).value,'%Y/%m/%d-%H:%M:%S-%f'))
    #print(datetime.datetime.strptime(sh.cell(14, 3).value,'%Y/%m/%d-%H:%M:%S-%f'))
    #print(deltaSecond(datetime.datetime.strptime(sh.cell(1, 3).value,'%Y/%m/%d-%H:%M:%S-%f'),datetime.datetime.strptime(sh.cell(14, 3).value,'%Y/%m/%d-%H:%M:%S-%f')))
    collect1=dict()
    reader.getBasicInfo(collect1)
    #print('xlsx', datetime.datetime(*tuple(map(lambda x: int(x.lstrip('0')),sh.cell(1, 3).value.split('-')[0].split('/')))))
    #print(collect1[sh.cell(0, 3).value])

    doc = Document()

    doc.styles['Normal'].font.name = u'標楷體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    doc.styles['Normal'].font.bold=True
    doc.styles['Normal'].font.size=Pt(12)

    ###############解決原生bug####################
    styles = doc.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    #new_heading_style.base_style = styles['Heading 1']

    font = new_heading_style.font
    font.bold=True
    font.name = u'標楷體'
    font.size = Pt(16)

    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')


    #print(isinstance(doc.styles['Heading 1'],p))
    #print(type(new_heading_style))
    #print(dir(new_heading_style.element))
    #doc.add_heading('D-KEFS執行功能測驗回饋單')
    #doc.add_heading('別師', level=1)
    paragraph = doc.add_paragraph('D-KEFS執行功能測驗回饋單', style='New Heading')
    paragraph.alignment = 1

    doc.add_paragraph('基本資料:')


    table = doc.add_table(rows = 3, cols = 6, style='TableGrid') 


    hdr_cells1 = table.rows[0].cells
    hdr_cells1[0].text = '編號:'
    hdr_cells1[1].text = str(int(collect1['ID']))
    hdr_cells1[2].text = '性別:'
    hdr_cells1[3].text = collect1['Gender']
    hdr_cells1[4].text = '測驗時間:'
    hdr_cells1[5].text = str(collect1['CreatTime'].strftime("%Y/%m/%d"))

    hdr_cells2 = table.rows[1].cells
    hdr_cells2[0].text = '姓名:'
    hdr_cells2[1].text = collect1['Name']
    hdr_cells2[2].text = '慣用手:'
    hdr_cells2[3].text = collect1['HabitualHand']
    hdr_cells2[4].text = '出生年月日:'
    hdr_cells2[5].text = str(collect1['DateOfBirth'].strftime("%Y/%m/%d"))

    hdr_cells3 = table.rows[2].cells
    hdr_cells3[0].text = '年齡:'
    hdr_cells3[1].text = str(collect1['CreatTime'].year-collect1['DateOfBirth'].year)
    hdr_cells3[2].text = '教育年:'
    hdr_cells3[3].text = str(int(collect1['HighestLevelOfEducation']))
    hdr_cells3[4].text = '學校:'





    doc.add_paragraph('')
    paragraph = doc.add_paragraph('軌跡標示測驗', style='New Heading')
    paragraph.alignment = 1
    doc.add_paragraph('基本測量:')
    table = doc.add_table(rows = 4, cols = 6, style='TableGrid') 
    hdr_cells1 = table.rows[0].cells
    hdr_cells1[1].text = '情境一：\n視覺掃描'
    hdr_cells1[2].text = '情境二：\n圓形序列'
    hdr_cells1[3].text = '情境三：\n六邊形序列'
    hdr_cells1[4].text = '情境四：\n圓形六邊形轉換'
    hdr_cells1[5].text = '情境五：\n動作速度'

    data={}
    reader.getData(data,'Task1')
    reader.getData( data,'Task2')
    reader.getData( data,'Task3')
    reader.getData( data,'Task4')
    reader.getData( data,'Task5')
    print('data', data)






    hdr_cells2=table.rows[1].cells
    hdr_cells2[0].text = '原始\n分數'
    hdr_cells2[1].text = str(data['Task1'])
    hdr_cells2[2].text = str(data['Task2'])
    hdr_cells2[3].text = str(data['Task3'])
    hdr_cells2[4].text = str(data['Task4'])
    hdr_cells2[5].text = str(data['Task5'])


    hdr_cells3=table.rows[2].cells
    hdr_cells3[0].text = '量尺\n分數'
    score2={}
    score2["Task1"]=newScore(data["Task1"], 20.5, 7.25, 10, 3)
    hdr_cells3[1].text =f'{score2["Task1"]:.3f}'
    score2['Task2']=newScore(data['Task2'], 29.5, 11, 10, 3)
    hdr_cells3[2].text =f"{score2['Task2']:.3f}"
    score2['Task3']=newScore(data['Task3'], 29, 10, 10, 3)
    hdr_cells3[3].text =f"{score2['Task3']:.3f}"
    score2['Task4']=newScore(data['Task4'], 69, 28, 10, 3)
    hdr_cells3[4].text =f"{score2['Task4']:.3f}"
    score2['Task5']=newScore(data['Task5'], 31, 16, 10, 3)
    hdr_cells3[5].text =f"{score2['Task5']:.3f}"



    hdr_cells4=table.rows[3].cells
    hdr_cells4[0].text = 'PR值'
    hdr_cells4[1].text =prValue(data["Task1"], 20.5, 7.25)
    hdr_cells4[2].text =prValue(data["Task2"], 29.5, 11)
    hdr_cells4[3].text =prValue(data["Task3"], 29, 10)
    hdr_cells4[4].text =prValue(data["Task4"], 69, 28)
    hdr_cells4[5].text =prValue(data["Task5"], 31, 16)




    doc.add_paragraph('')
    doc.add_paragraph('衍生測量:')
    table = doc.add_table(rows = 4, cols = 7, style='TableGrid')
    hdr_cells1 = table.rows[0].cells
    hdr_cells1[1].text = '圓形序列＋六邊形序列'
    hdr_cells1[2].text = '圓形六邊形轉換 - 視覺掃描'
    hdr_cells1[3].text = '圓形六邊形轉換 - 圓形序列'
    hdr_cells1[4].text = '圓形六邊形轉換 - 六邊形序列'
    hdr_cells1[5].text = '圓形六邊形轉換- 圓形序列＋六邊形序列'
    hdr_cells1[6].text = '圓形六邊形轉換 - 動作速度'



    combine={}
    hdr_cells2 = table.rows[1].cells
    i=0
    hdr_cells2[i].text = '計分\n結果'
    i+=1
    combine[i]=score2['Task2']+score2['Task3']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    i+=1
    combine[i]=score2['Task1']-score2['Task2']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    i+=1
    combine[i]=score2['Task4']-score2['Task2']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    i+=1
    combine[i]=score2['Task4']-score2['Task3']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    i+=1
    combine[i]=score2['Task4']-score2['Task2']+score2['Task3']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    i+=1
    combine[i]=score2['Task4']-score2['Task5']
    hdr_cells2[i].text = f"{combine[i]:.3f}"
    



    new_combine={}
    hdr_cells3 = table.rows[2].cells
    i=0
    hdr_cells3[i].text = '量尺\n分數'
    i+=1
    new_combine[i]=newScore(combine[i], 19.5, 5.25, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'
    i+=1
    new_combine[i]=newScore(combine[i], 0, 3.75, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'
    i+=1
    new_combine[i]=newScore(combine[i], 0, 3, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'
    i+=1
    new_combine[i]=newScore(combine[i], 0, 3, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'
    i+=1
    new_combine[i]=newScore(combine[i], 0, 3, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'
    i+=1
    new_combine[i]=newScore(combine[i], 0, 3, 10, 3)
    hdr_cells3[i].text = f'{new_combine[i]:.3f}'

    #print('new combine: ', new_combine)



    hdr_cells4 = table.rows[3].cells
    i=0
    hdr_cells4[i].text = 'PR值'
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 19.5, 5.25)
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 0, 3.75)
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 0, 3)
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 0, 3)
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 0, 3)
    i+=1
    hdr_cells4[i].text = prValue(combine[i], 0, 3)
    


    doc.add_paragraph('')
    doc.add_paragraph('選擇性測量：錯誤分析')
    table = doc.add_table(rows = 7, cols = 6, style='TableGrid') 
    hdr_cells1 = table.rows[0].cells
    hdr_cells1[1].text = '情境一：\n視覺掃描'
    hdr_cells1[2].text = '情境二：\n圓形序列'
    hdr_cells1[3].text = '情境三：\n六邊形序列'
    hdr_cells1[4].text = '情境四：\n圓形六邊形轉換'
    hdr_cells1[5].text = '情境五：\n動作速度'

    table.rows[1].cells[0].text = '疏忽錯誤'
    table.rows[2].cells[0].text = '任務錯誤'
    table.rows[3].cells[0].text = '序列錯誤'
    table.rows[4].cells[0].text = '不正確反應錯誤'
    table.rows[5].cells[0].text = '時間中止錯誤'
    table.rows[6].cells[0].text = '錯誤總數'
    doc.add_paragraph('*原始分數/累積百分位數')
    doc.add_paragraph('')
    doc.add_paragraph('說明:')

    #print(sh.nrows)
    #print(dir(doc.styles['Normal'].font))
    file=file.split('.')[0]
    

    doc.save(destFolder+'/word/'+file+'.docx')


    
    #doc to pdf
    wdFormatPDF = 17

    in_file = destFolder+'\\word\\'+file+'.docx'
    out_file = destFolder+'\\pdf\\'+file+'.pdf'

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

'''  
if not os.path.exists('word'):
        os.makedirs('word')
if not os.path.exists('pdf'):
    os.makedirs('pdf')
for file in glob.glob("*.xlsx"):
    reader=Reader(file)
    excel2doc(reader, file)
'''





try:
    root = Tk()
    root.title("excel to pdf...")
    windowWidth = root.winfo_reqwidth()
    windowHeight = root.winfo_reqheight()
    positionRight = int(root.winfo_screenwidth()/2 - windowWidth/2)
    positionDown = int(root.winfo_screenheight()/2 - windowHeight/2)
    root.geometry(f"+{positionRight}+{positionDown}")

    pb = Progressbar(root,length=200,mode="determinate",orient=HORIZONTAL)
    pb.pack(padx=10,pady=10)
    
    
    

    
    if not os.path.exists('word'):
        os.makedirs('word')
    if not os.path.exists('pdf'):
        os.makedirs('pdf')
    
    pb["maximum"] = len(glob.glob("*.xlsx"))
    pb["value"] = 0
    
    i=0
    folder_selected = filedialog.askdirectory()
    for file in glob.glob("*.xlsx"):
        reader=Reader(file)
        excel2doc(reader, file,folder_selected)
        i+=1
        pb["value"] = i
        root.update()

        
    messagebox.showinfo(root,"成功將所有 Excel 輸出成 pdf")
    root.destroy()
    root.mainloop()
except ZeroDivisionError as e:
    print('zeroDivisionError')
'''
except Exception as e:
    print('錯誤!!!!')
    print(type(e).__module__, type(e).__qualname__)
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉\n\n錯誤訊息: "+str(e))
    
    root.destroy()
    

except:
    print('錯誤!!!!')
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉")
    root.destroy()


'''
    


