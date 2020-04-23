from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import datetime
from docx.shared import  Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import comtypes.client
from scipy.stats import norm
import glob
from tkinter import *
from tkinter.ttk import *
from tkinter import messagebox
from tkinter import filedialog
import traceback
import csv

def callback_error(*args):
    # Build the error message
    message = 'Generic error:\n\n'
    message += traceback.format_exc()

    # Also log the error to a file
    # TODO
    # Show the error to the user
    print(message)

    # Exit the program immediately
    exit()


def chooseDest():
    global dest_folder
    old_folder=dest_folder
    dest_folder = filedialog.askdirectory()
    if not dest_folder:
        dest_folder =old_folder
    labeldest["text"] = '輸出資料夾:  '+dest_folder






def go(dest_folder, source_folder):
    i=0
    '''
    if not os.path.exists('word'):
        os.makedirs('word')
    if not os.path.exists('pdf'):
        os.makedirs('pdf')
    '''
    if not os.path.isdir(dest_folder+'/word'):
        os.makedirs(dest_folder+'/word')
    if not os.path.isdir(dest_folder+'/pdf'):
        os.makedirs(dest_folder+'/pdf')
    

    filelist = [ f for f in os.listdir(dest_folder+'/pdf') if f.endswith(".pdf") ]
    for f in filelist:
        os.remove(os.path.join(dest_folder+'/pdf', f))

    for file in glob.glob(source_folder+"\\*.csv"):
        #print('file',file)
        #print('source: ',source_folder)
        #file=os.path.basename(file)
        with open(file,newline='') as csvfile:
            rows = list(csv.reader(csvfile))
            reader=Reader(rows)
            writer=Writer()
            excel2doc(reader,writer, file,dest_folder)
            i+=1
            pb["value"] = i
            root.update()

        
    messagebox.showinfo(root,"成功將所有 Excel 輸出成 pdf")
    root.destroy()


class Writer():
    def __init__(self):
        self.doc = Document()
        self.doc.styles['Normal'].font.name = u'標楷體'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        self.doc.styles['Normal'].font.bold=True
        self.doc.styles['Normal'].font.size=Pt(12)
        styles = self.doc.styles
        new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)

        font = new_heading_style.font
        font.bold=True
        font.name = u'標楷體'
        font.size = Pt(16)

        new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    def add_title(self,title):
        paragraph = self.doc.add_paragraph(title, style='New Heading')
        paragraph.alignment = 1

    def add_paragraph(self,context):
        self.doc.add_paragraph(context)
    def add_table(self,size,context):
        table = self.doc.add_table(rows = size[0], cols = size[1], style='TableGrid')
        for index, item in enumerate(context):
            for index2, item2 in enumerate(item):
                table.cell(index,index2).text=str(item2)
    
    def save(self,path):
        self.doc.save(path)



class Reader():
    source_folder=os.getcwd()
    #source_folder can only be used in Reader.source_folder and it cannot be used in any instance
    def __init__(self, rows):
        self.rows=rows


    @staticmethod
    def chooseSource():
        old_folder=Reader.source_folder
        Reader.source_folder = filedialog.askdirectory()
        if not Reader.source_folder:
            Reader.source_folder =old_folder
        labelsource["text"] = '輸入資料夾:  '+Reader.source_folder
        pb["maximum"] = len(glob.glob(Reader.source_folder+"\\*.csv"))
    
    def getData(self, task):
        rows=self.rows
        i=0
        while True:
            if rows[i][0]==task:
                i+=1
                while rows[i][0]=='':
                    if rows[i][1]=='Complete_Time':
                        return float(f'{float(rows[i][2])/1000:.3f}')
                        break
                    i+=1
                break
            i+=1




    def getBasicInfo(self, collect1,):
        rows=self.rows
        i=0
        while True:
            if rows[i][0]=='個人資料':
                collect1[rows[i][1]]=rows[i][2]
                i+=1
                while rows[i][0]=='':
                    if rows[i][1]=='DateOfBirth':
                        collect1[rows[i][1]]=datetime.datetime(*map(int,rows[i][2].split('/')))
                    else:
                        collect1[rows[i][1]]=rows[i][2]
                    i+=1
                break
            i+=1
        collect1[rows[0][3]]=datetime.datetime.strptime(rows[1][3],'%Y/%m/%d-%H:%M:%S-%f')
        '''
        for count in range(sh.nrows):
            if sh.cell(count, 0).value=='個人資料':
                collect1[sh.cell(count, 1).value]=sh.cell(count, 2).value
                count+=1
                while  sh.cell(count, 0).value=='':
                    if sh.cell(count, 1).value=='DateOfBirth':
                        collect1[sh.cell(count, 1).value]=datetime.datetime(*xlrd.xldate_as_tuple(sh.cell(count, 2).value, book.datemode))
                    else:
                        collect1[sh.cell(count, 1).value]=sh.cell(count, 2).value
                    count+=1
                break

        #print(collect1)
        collect1[sh.cell(0, 3).value]=datetime.datetime.strptime(sh.cell(1, 3).value,'%Y/%m/%d-%H:%M:%S-%f')
        '''





def prValue(data,mean,std, norm=norm):
    if norm.cdf(data,mean,std)*100<1:
        return str(1)
    else:
        return str(int(norm.cdf(data,mean,std)*100))









def newScore(data, mean,std, new_mean, new_std):
    return float(f'{(data-mean)/std*new_std+new_mean:.3f}')



def excel2doc(reader,writer, file, destFolder):
    collect1=dict()
    reader.getBasicInfo(collect1)
    print('collect1',collect1)
    writer.add_title('D-KEFS執行功能測驗回饋單')

    
    
    writer.add_paragraph('基本資料:')
    ##table1
    tableList=[
    ['編號:',str(int(collect1['ID'])),                                    '性別:'  ,collect1['Gender'],       '測驗時間:',   str(collect1['CreatTime'].strftime("%Y/%m/%d"))  ],
    ['姓名:',collect1['Name'],                                            '慣用手:',collect1['HabitualHand'],  '出生年月日:', str(collect1['DateOfBirth'].strftime("%Y/%m/%d"))],
    ['年齡:',str(collect1['CreatTime'].year-collect1['DateOfBirth'].year),'教育年:',collect1['HighestLevelOfEducation'], '學校:']
    ]
    writer.add_table([3,6],tableList)
    writer.add_paragraph('')
    writer.add_title('軌跡標示測驗')
    writer.add_paragraph('基本測量:')


    ##table2
    i1=['',         '情境一：\n視覺掃描',                                '情境二：\n圓形序列',              '情境三：\n六邊形序列',          '情境四：\n圓形六邊形轉換',     '情境五：\n動作速度']
    i2=['原始\n分數',reader.getData('Task1'),                           reader.getData('Task2'),          reader.getData('Task3'),        reader.getData('Task4'),       reader.getData('Task5')]
    print(i2)
    i3=['量尺\n分數',newScore(i2[1], 20.5, 7.25, 10, 3),newScore(i2[2], 29.5, 11, 10, 3), newScore(i2[3], 29, 10, 10, 3),newScore(i2[4], 69, 28, 10, 3),newScore(i2[5], 31, 16, 10, 3)]
    i4=['PR值',     prValue(i2[1], 20.5, 7.25),                 prValue(i2[2], 29.5, 11),prValue(i2[3], 29, 10),prValue(i2[4], 69, 28), prValue(i2[5], 31, 16)]
    tableList2=[i1,i2,i3,i4]
    writer.add_table([4,6],tableList2)


    ##table3
    table2=tableList2[2]
    i1=['',         '圓形序列＋六邊形序列',            '圓形六邊形轉換 - 視覺掃描',    '圓形六邊形轉換 - 圓形序列','圓形六邊形轉換 - 六邊形序列','圓形六邊形轉換- 圓形序列＋六邊形序列','圓形六邊形轉換 - 動作速度']
    i2=['計分\n結果',float(f'{table2[2]+table2[3]:.3}'),   float(f'{table2[1]-table2[2]:.3}'),        float(f'{table2[4]-table2[2]:.3}'),      float(f'{table2[4]-table2[3]:.3}'),       float(f'{table2[4]-table2[2]+table2[3]:.3}'),                float(f'{table2[4]-table2[5]:.3}')]
    i3=['量尺\n分數',newScore(i2[1], 19.5, 5.25, 10, 3), newScore(i2[2], 0, 3.75, 10, 3),newScore(i2[3], 0, 3, 10, 3), newScore(i2[4], 0, 3, 10, 3),    newScore(i2[5], 0, 3, 10, 3),  newScore(i2[6], 0, 3, 10, 3)]
    i4=['PR值',      prValue(i2[1], 19.5, 5.25),prValue(i2[2], 0, 3.75),prValue(i2[3], 0, 3), prValue(i2[4], 0, 3),  prValue(i2[5], 0, 3),           prValue(i2[6], 0, 3)]
    tableList3=[i1,i2,i3,i4]
    writer.add_paragraph('')
    writer.add_paragraph('衍生測量:')

    writer.add_table([4,7],tableList3)
    
    writer.add_paragraph('')
    writer.add_paragraph('選擇性測量：錯誤分析')
    ##table4
    tableList4=[
    ['',        '情境一：\n視覺掃描','情境二：\n圓形序列','情境三：\n六邊形序列','情境四：\n圓形六邊形轉換','情境五：\n動作速度'],
    ['疏忽錯誤'],
    ['任務錯誤'],
    ['序列錯誤'],
    ['不正確反應錯誤'],
    ['時間中止錯誤'],
    ['錯誤總數']
    ]
    writer.add_table([7,6],tableList4)
    writer.add_paragraph('*原始分數/累積百分位數')
    writer.add_paragraph('')
    writer.add_paragraph('說明:')
    

    #print(sh.nrows)
    #print(dir(doc.styles['Normal'].font))
    
    file=os.path.basename(file)
    file=file.split('.')[0]
    writer.save(destFolder+'/word/'+file+'.docx')
    #doc to pdf
    wdFormatPDF = 17
    in_file = destFolder+'/word/'+file+'.docx'
    out_file = destFolder+'/pdf/'+file+'.pdf'
    #print('infile',in_file)
    #print('outfile',out_file)
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    






try:
    
    root = Tk()
    Tk.report_callback_exception = callback_error 
    root.title("excel to pdf...")
    windowWidth = root.winfo_reqwidth()
    windowHeight = root.winfo_reqheight()
    positionRight = int(root.winfo_screenwidth()/2 - windowWidth/2)
    positionDown = int(root.winfo_screenheight()/2 - windowHeight/2)
    root.geometry(f"+{positionRight}+{positionDown}")

    Reader.source_folder=os.getcwd()
    dest_folder =os.getcwd()
    labeldest = Label(root)                 # 標籤內容             
    btndest = Button(root,text="選擇路徑",command=lambda: chooseDest())
    btnGo = Button(root,text="執行",command=lambda: go(dest_folder,Reader.source_folder))
    labeldest.grid(row=0,column=0,padx=10,pady=10)       
    btndest.grid(row=0,column=1,padx=10,pady=10, sticky="w")   
    


    labelsource = Label(root) 
    btnsource = Button(root,text="選擇路徑",command=lambda: Reader.chooseSource())
    labelsource.grid(row=1,column=0,padx=10,pady=10)         
                  
    btnsource.grid(row=1,column=1,padx=10,pady=10, sticky="w")  

    Separator(root,orient=HORIZONTAL).grid(row=2, columnspan=2, sticky="ew")

    btnGo.grid(row=3,column=0,padx=10,pady=10, sticky="w")   
    pb = Progressbar(root,length=200,mode="determinate",orient=HORIZONTAL)
    pb.grid(row=3,column=1,padx=10,pady=10, sticky="w")
    labeldest["text"] = '輸出資料夾:  '+dest_folder
    labelsource['text']='讀取資料夾   '+Reader.source_folder
    
    pb["maximum"] = len(glob.glob(Reader.source_folder+"\\*.csv"))
    #glob.glob('*/*.csv')
    pb["value"] = 0
    
    root.mainloop()
except Exception as e:
    print('錯誤!!!!')
    print(type(e).__module__, type(e).__qualname__)
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉\n\n錯誤訊息: "+str(e))
    
    root.destroy()
except:
    print('錯誤!!!!')
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉")
    root.destroy()


    


