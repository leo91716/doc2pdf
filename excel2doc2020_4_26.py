from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import datetime
from docx.shared import  Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import comtypes.client
from scipy.stats import norm
import glob
from tkinter import *
from tkinter.ttk import *
from tkinter import messagebox
from tkinter import filedialog
import traceback
import csv
from tkinter.ttk import Combobox


def callback_error(*args):
    # Build the error message
    message = 'Generic error:\n\n'
    message += traceback.format_exc()
    print(message)

    # Exit the program immediately
    exit()


def chooseDest():
    global dest_folder
    old_folder=dest_folder
    dest_folder = filedialog.askdirectory()
    if not dest_folder:
        dest_folder =old_folder
    labeldest["text"] = '輸出資料夾:  '+dest_folder

class Primary():
    def go(self,dest_folder, source_folder,exper):
        expList={'軌跡標示測驗':'TMTest','設計流暢測驗':'DFTest'}
        i=0
        if not os.path.isdir(dest_folder+'/word'):
            os.makedirs(dest_folder+'/word')
        if not os.path.isdir(dest_folder+'/pdf'):
            os.makedirs(dest_folder+'/pdf')
        

        filelist = [ f for f in os.listdir(dest_folder+'/pdf') if f.endswith(".pdf") ]
        for f in filelist:
            os.remove(os.path.join(dest_folder+'/pdf', f))
        if len(glob.glob(source_folder+"\\"+expList[exper]+"_*.csv"))==0:
            messagebox.showwarning(root," \n 失敗: 找不到要讀的檔案")
            exit()
        pb["maximum"] = len(glob.glob(source_folder+"\\"+expList[exper]+"_*.csv"))
        for file in glob.glob(source_folder+"\\"+expList[exper]+"_*.csv"):
            with open(file,newline='') as csvfile:
                rows = list(csv.reader(csvfile))
                reader=Reader(rows)
                if exper=='軌跡標示測驗':
                    writer=Writer_Track(reader)
                else:
                    writer=Writer_Fluent(reader)
                self.excel2doc(writer, file,dest_folder)
                i+=1
                pb["value"] = i
                root.update()

            
        messagebox.showinfo(root,"成功將所有 Excel 輸出成 pdf")
        root.destroy()
    def excel2doc(self,writer, file, destFolder):
        writer.write_all()
        self.file=file
        self.destFolder=destFolder
        self.writer=writer
        self.save()
    def save(self):
        file=self.file
        writer=self.writer
        destFolder=self.destFolder
        file=os.path.basename(file)
        file=file.split('.')[0]
        writer.save(destFolder+'/word/'+file+'.docx')
        #doc to pdf
        wdFormatPDF = 17
        in_file = destFolder+'/word/'+file+'.docx'
        out_file = destFolder+'/pdf/'+file+'.pdf'
        #print('infile',in_file)
        #print('outfile',out_file)
        self.word = comtypes.client.CreateObject('Word.Application')
        self.doc = self.word.Documents.Open(in_file)
        self.doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        self.doc.Close()
        self.word.Quit()


class Writer_Track():
    def __init__(self,reader):
        self.reader=reader
        self.doc = Document()
        self.doc.styles['Normal'].font.name = u'標楷體'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        self.doc.styles['Normal'].font.bold=True
        self.doc.styles['Normal'].font.size=Pt(12)
        styles = self.doc.styles
        new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)

        font = new_heading_style.font
        font.bold=True
        font.name = u'標楷體'
        font.size = Pt(16)

        new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    def add_title(self,title):
        paragraph = self.doc.add_paragraph(title, style='New Heading')
        paragraph.alignment = 1

    def add_paragraph(self,context):
        self.doc.add_paragraph(context)
    def add_table(self,size,context, merge=[]):
        table = self.doc.add_table(rows = size[0], cols = size[1], style='TableGrid')
        for index, item in enumerate(context):
            for index2, item2 in enumerate(item):
                if type(item2)==float:
                    table.cell(index,index2).text=f'{item2:.3f}'
                else:
                    table.cell(index,index2).text=str(item2)
                
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = 1
        if merge:
            a=table.cell(*merge[0])
            b=table.cell(*merge[1])
            a.merge(b)
    
    def save(self,path):
        self.doc.save(path)

    def getBasicInfo(self):
        collect1=dict()
        self.reader.getBasicInfo(collect1)
        print('collect1',collect1)
        self.add_title('D-KEFS執行功能測驗回饋單')
        self.add_paragraph('基本資料:')
        tableList=[
        ['編號:',str(int(collect1['ID'])),                                    '性別:'  ,collect1['Gender'],       '測驗時間:',   str(collect1['CreatTime'].strftime("%Y/%m/%d"))  ],
        ['姓名:',collect1['Name'],                                            '慣用手:',collect1['HabitualHand'],  '出生年月日:', str(collect1['DateOfBirth'].strftime("%Y/%m/%d"))],
        ['年齡:',str(collect1['CreatTime'].year-collect1['DateOfBirth'].year),'教育年:',collect1['HighestLevelOfEducation'], '學校:']
        ]
        self.add_table([3,6],tableList)
        

    def getBasicMeasure(self):
        self.add_paragraph('')
        self.add_title('軌跡標示測驗')
        self.add_paragraph('基本測量:')
        reader=self.reader
        i1=['',         '情境一：\n視覺掃描',                                '情境二：\n圓形序列',              '情境三：\n六邊形序列',          '情境四：\n圓形六邊形轉換',     '情境五：\n動作速度']
        i2=['原始\n分數',reader.getData('Task1','Complete_Time'),                           reader.getData('Task2','Complete_Time'),          reader.getData('Task3','Complete_Time'),        reader.getData('Task4','Complete_Time'),       reader.getData('Task5','Complete_Time')]
        print(i2)
        i3=['量尺\n分數',newScore(i2[1], 20.5, 7.25, 10, 3),newScore(i2[2], 29.5, 11, 10, 3), newScore(i2[3], 29, 10, 10, 3),newScore(i2[4], 69, 28, 10, 3),newScore(i2[5], 31, 16, 10, 3)]
        i4=['PR值',     prValue(i2[1], 20.5, 7.25),                 prValue(i2[2], 29.5, 11),prValue(i2[3], 29, 10),prValue(i2[4], 69, 28), prValue(i2[5], 31, 16)]
        self.tableList2=[i1,i2,i3,i4]
        self.add_table([4,6],self.tableList2)

    def getMoreMeasure(self):
        ##table3
        table2=self.tableList2[2]
        i1=['',         '圓形序列＋六邊形序列',            '圓形六邊形轉換 - 視覺掃描',    '圓形六邊形轉換 - 圓形序列','圓形六邊形轉換 - 六邊形序列','圓形六邊形轉換- 圓形序列＋六邊形序列','圓形六邊形轉換 - 動作速度']
        i2=['計分\n結果',table2[2]+table2[3],   table2[1]-table2[2],        table2[4]-table2[2],      table2[4]-table2[3],       table2[4]-table2[2]-table2[3],                table2[4]-table2[5]]
        i3=['量尺\n分數',newScore(i2[1], 19.5, 5.25, 10, 3), newScore(i2[2], 0, 3.75, 10, 3),newScore(i2[3], 0, 3, 10, 3), newScore(i2[4], 0, 3, 10, 3),    newScore(i2[5], 0, 3, 10, 3),  newScore(i2[6], 0, 3, 10, 3)]
        i4=['PR值',      prValue(i2[1], 19.5, 5.25),prValue(i2[2], 0, 3.75),prValue(i2[3], 0, 3), prValue(i2[4], 0, 3),  prValue(i2[5], 0, 3),           prValue(i2[6], 0, 3)]
        tableList3=[i1,i2,i3,i4]
        self.add_paragraph('')
        self.add_paragraph('衍生測量:')
        self.add_table([4,7],tableList3)

    def getOptionalTable(self):
        self.doc.add_page_break()
        self.add_paragraph('選擇性測量：錯誤分析')
        ##table4
        tableList4=[
        ['',        '情境一：\n視覺掃描','情境二：\n圓形序列','情境三：\n六邊形序列','情境四：\n圓形六邊形轉換','情境五：\n動作速度'],
        ['疏忽錯誤'],
        ['任務錯誤'],
        ['序列錯誤'],
        ['不正確反應錯誤'],
        ['時間中止錯誤'],
        ['錯誤總數']
        ]
        self.add_table([7,6],tableList4)
        self.add_paragraph('*原始分數/累積百分位數')
    def write_all(self):
        self.getBasicInfo()
        self.getBasicMeasure()
        self.getMoreMeasure()
        self.getOptionalTable()
        self.add_paragraph('')
        self.add_paragraph('說明:')

class Writer_Fluent(Writer_Track):  #newScore(data, mean,std, new_mean, new_std)
    def getBasicMeasure(self):
        self.add_paragraph('')
        self.add_title('設計流暢性測驗')
        self.add_paragraph('基本測量:')
        reader=self.reader
        getDataArg=('Examing','Total_CorrectDesign',False)
        i0=['','總正確數']
        i1=['',         '情境1\n黑點相連',                                '情境2\n白點相連',             '情境3 黑點和白點互相轉換','設計流暢性\n總和']
        i2=['原始分數',int(reader.getData(*getDataArg,'情境1_黑點相連')),                           int(reader.getData(*getDataArg,'情境2_白點相連')),          int(reader.getData(*getDataArg,'情境3_黑點和白點互相轉換'))]
        i2end=[i2[1]+i2[2]+i2[3]]
        i2=i2+i2end
        #print(i2)
        i3=['量尺\n分數',newScore(i2[1], 10, 3.75, 10, 3),newScore(i2[2], 11, 4, 10, 3), newScore(i2[3], 8, 3, 10, 3),newScore(i2[4], 29, 7.5, 10, 3)]
        #i4=['PR值',     prValue(i2[1], 20.5, 7.25),                 prValue(i2[2], 29.5, 11),prValue(i2[3], 29, 10),prValue(i2[4], 69, 28), prValue(i2[5], 31, 16)]
        self.tableList2=[i0,i1,i2,i3]
        merge=[(0,1),(0,4)]
        self.add_table([5,5],self.tableList2,merge)
    
    def getMoreMeasure(self):
        table2=self.tableList2[3]
        i1=['',         '實心點連結＋空心點連結\n組合總正確數','轉換 - 實心點連結＋空心點連結']
        i2=['計分結果',table2[1]+table2[2],   table2[3]-table2[1]-table2[2]]
        i3=['量尺分數',newScore(i2[1], 19.5, 6, 10, 3), newScore(i2[2], 0, 3, 10, 3)]
        i4=['PR值']
        tableList3=[i1,i2,i3,i4]
        self.add_paragraph('')
        self.add_paragraph('衍生測量:')
        self.add_table([4,3],tableList3)

    def getOptionalTable(self): #getData(self, task,item,thousand=True,doubleCheck=None)
        self.add_paragraph('')
        self.add_paragraph('選擇性測量')
        getData=self.reader.getData
        table2=self.tableList2[2]
        getDataArg1=['Examing','Total_UnCorrectDesign',False]
        getDataArg2=['Examing','Total_RepeatDesign',False]
        getDataArg3=['Examing','Total_TryDesign',False]
        ##table4
        
        i1=['',        '不正確設計數','重複設計','嘗試設計總數','正確設計\n百分比']

        i2=['原始總分',getData(*getDataArg1,'情境1_黑點相連')+getData(*getDataArg1,'情境2_白點相連')+getData(*getDataArg1,'情境3_黑點和白點互相轉換'),
        getData(*getDataArg2,'情境1_黑點相連')+getData(*getDataArg2,'情境2_白點相連')+getData(*getDataArg2,'情境3_黑點和白點互相轉換'),
        getData(*getDataArg3,'情境1_黑點相連')+getData(*getDataArg3,'情境2_白點相連')+getData(*getDataArg3,'情境3_黑點和白點互相轉換')]
        i2End=[table2[4]/i2[3]]
        i2=i2+i2End
        i3=['量尺分數']
        i4=['PR值']
        tableList4=[i1,i2,i3,i4]
        self.add_table([4,5],tableList4)


class Reader():
    source_folder=os.getcwd()
    #source_folder can only be used in Reader.source_folder and it cannot be used in any instance
    def __init__(self, rows):
        self.rows=rows


    @staticmethod
    def chooseSource():
        old_folder=Reader.source_folder
        Reader.source_folder = filedialog.askdirectory()
        if not Reader.source_folder:
            Reader.source_folder =old_folder
        labelsource["text"] = '輸入資料夾:  '+Reader.source_folder
        #pb["maximum"] = len(glob.glob(Reader.source_folder+"\\*.csv"))
    
    def getData(self, task,item,thousand=True,doubleCheck=None):
        rows=self.rows
        i=0
        while i<len(rows):
            if rows[i][0]==task:
                if doubleCheck==None or rows[i][2]==doubleCheck:
                    i+=1
                    while rows[i][0]=='':
                        if rows[i][1]==item:    #'Complete_Time'
                            if thousand:
                                thousandNumber=1000
                            else:
                                thousandNumber=1
                            return float(rows[i][2])/thousandNumber
                            break
                        i+=1
                    break
            i+=1




    def getBasicInfo(self, collect1):
        rows=self.rows
        i=0
        while True:
            if rows[i][0]=='個人資料':
                collect1[rows[i][1]]=rows[i][2]
                i+=1
                while rows[i][0]=='':
                    if rows[i][1]=='DateOfBirth':
                        collect1[rows[i][1]]=datetime.datetime(*map(int,rows[i][2].split('/')))
                    else:
                        collect1[rows[i][1]]=rows[i][2]
                    i+=1
                break
            i+=1
        collect1[rows[0][3]]=datetime.datetime.strptime(rows[1][3],'%Y/%m/%d-%H:%M:%S-%f')




def prValue(data,mean,std, norm=norm):
    if norm.cdf(data,mean,std)*100<1:
        return str(1)
    else:
        return str(int(norm.cdf(data,mean,std)*100))


def newScore(data, mean,std, new_mean, new_std):
    return (data-mean)/std*new_std+new_mean



    
    






try:
    
    root = Tk()
    Tk.report_callback_exception = callback_error 
    root.title("excel to pdf...")
    windowWidth = root.winfo_reqwidth()
    windowHeight = root.winfo_reqheight()
    positionRight = int(root.winfo_screenwidth()/2 - windowWidth/2)
    positionDown = int(root.winfo_screenheight()/2 - windowHeight/2)
    root.geometry(f"+{positionRight}+{positionDown}")
    primary=Primary()
    Reader.source_folder=os.getcwd()
    dest_folder =os.getcwd()
    labeldest = Label(root)                 # 標籤內容             
    btndest = Button(root,text="選擇路徑",command=lambda: chooseDest())
    btnGo = Button(root,text="執行",command=lambda: primary.go(dest_folder,Reader.source_folder,expVar.get()))
    labeldest.grid(row=0,column=0,padx=10,pady=5, sticky="w")       
    btndest.grid(row=0,column=1,padx=10,pady=5, sticky="w")   
    


    labelsource = Label(root) 
    btnsource = Button(root,text="選擇路徑",command=lambda: Reader.chooseSource())
    labelsource.grid(row=1,column=0,padx=10,pady=5, sticky="w")         
    btnsource.grid(row=1,column=1,padx=10,pady=5, sticky="w")

    labelExp = Label(root,text='實驗:') 
    labelExp.grid(row=2,column=0,padx=10,pady=5, sticky="e")   
    expVar = StringVar()       
    cb = Combobox(root,textvariable=expVar,        # 建立Combobox
                value=("軌跡標示測驗","設計流暢測驗"))   
    cb.current(0)
    cb.grid(row=2,column=1,padx=10,pady=5, sticky="w")     



    Separator(root,orient=HORIZONTAL).grid(row=3, columnspan=2, sticky="ew")

    btnGo.grid(row=4,column=0,padx=10,pady=5, sticky="e")   
    pb = Progressbar(root,length=200,mode="determinate",orient=HORIZONTAL)
    pb.grid(row=4,column=1,padx=10,pady=5, sticky="w")
    labeldest["text"] = '輸出資料夾:  '+dest_folder
    labelsource['text']='讀取資料夾   '+Reader.source_folder
    
    pb["maximum"] = len(glob.glob(Reader.source_folder+"\\*.csv"))
    #glob.glob('*/*.csv')
    pb["value"] = 0
    
    root.mainloop()

except Exception as e:
    try:
        primary.save()
    except:
        print('resave error')
    print('錯誤!!!!')
    print(type(e).__module__, type(e).__qualname__)
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉\n\n錯誤訊息: "+str(e))
    
    root.destroy()
    
except:
    try:
        primary.save()
    except:
        print('resave error')
    print('錯誤!!!!')
    messagebox.showwarning(root," \n 失敗: 無法將所有 Excel 輸出成 pdf\n\n提醒:所有檔案必須關閉 ")
    root.destroy()



    


